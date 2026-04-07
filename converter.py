"""
converter.py — Pillow 12.x негізіндегі түрлендіру логикасы.

Жалғыз сурет:  convert_single(data, fmt, filename)
Көп сурет:     convert_multi(images, fmt, filename)   ← PDF / DOCX ғана
"""

from __future__ import annotations

import io
from PIL import Image, ImageOps
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from docx import Document
from docx.shared import Inches


# ──────────────────────────────────────────────
# Internal helpers
# ──────────────────────────────────────────────

def _open(data: bytes) -> Image.Image:
    img = Image.open(io.BytesIO(data))
    return ImageOps.exif_transpose(img)


def _to_rgb(img: Image.Image) -> Image.Image:
    if img.mode not in ("RGB",):
        return img.convert("RGB")
    return img


def _img_to_jpeg_buf(img: Image.Image) -> io.BytesIO:
    img = _to_rgb(img)
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=95)
    buf.seek(0)
    return buf


# ──────────────────────────────────────────────
# Single converters
# ──────────────────────────────────────────────

def _single_pdf(img: Image.Image) -> bytes:
    A4_W, A4_H, PAD = 595.0, 842.0, 20.0
    mw, mh = A4_W - PAD * 2, A4_H - PAD * 2
    scale = min(mw / img.width, mh / img.height)
    dw, dh = img.width * scale, img.height * scale
    x, y = (A4_W - dw) / 2, (A4_H - dh) / 2

    out = io.BytesIO()
    c = canvas.Canvas(out, pagesize=(A4_W, A4_H))
    c.drawImage(ImageReader(_img_to_jpeg_buf(img)), x, y, width=dw, height=dh)
    c.save()
    return out.getvalue()


def _single_docx(img: Image.Image) -> bytes:
    doc = Document()
    sec = doc.sections[0]
    page_w = sec.page_width - sec.left_margin - sec.right_margin

    img_buf = io.BytesIO()
    img.save(img_buf, format="PNG")
    img_buf.seek(0)

    doc.add_picture(img_buf, width=page_w)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _single_jpg(img: Image.Image) -> bytes:
    buf = io.BytesIO()
    _to_rgb(img).save(buf, format="JPEG", quality=95)
    return buf.getvalue()


def _single_png(img: Image.Image) -> bytes:
    buf = io.BytesIO()
    img.save(buf, format="PNG", optimize=True)
    return buf.getvalue()


# ──────────────────────────────────────────────
# Multi-image converters  (PDF / DOCX ғана)
# ──────────────────────────────────────────────

def _multi_pdf(images: list[Image.Image]) -> bytes:
    A4_W, A4_H, PAD = 595.0, 842.0, 20.0
    mw, mh = A4_W - PAD * 2, A4_H - PAD * 2

    out = io.BytesIO()
    c = canvas.Canvas(out, pagesize=(A4_W, A4_H))

    for idx, img in enumerate(images):
        if idx > 0:
            c.showPage()
        scale = min(mw / img.width, mh / img.height)
        dw, dh = img.width * scale, img.height * scale
        x, y = (A4_W - dw) / 2, (A4_H - dh) / 2
        c.drawImage(ImageReader(_img_to_jpeg_buf(img)), x, y, width=dw, height=dh)

    c.save()
    return out.getvalue()


def _multi_docx(images: list[Image.Image]) -> bytes:
    doc = Document()
    sec = doc.sections[0]
    page_w = sec.page_width - sec.left_margin - sec.right_margin

    for idx, img in enumerate(images):
        if idx > 0:
            # Жаңа бет
            doc.add_page_break()

        img_buf = io.BytesIO()
        img.save(img_buf, format="PNG")
        img_buf.seek(0)
        doc.add_picture(img_buf, width=page_w)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# ──────────────────────────────────────────────
# Public API
# ──────────────────────────────────────────────

def convert_single(data: bytes, fmt: str, filename: str) -> tuple[bytes, str]:
    """Бір суретті түрлендіру. (bytes, filename) қайтарады."""
    img = _open(data)
    dispatch = {
        "pdf":  (_single_pdf,  ".pdf"),
        "docx": (_single_docx, ".docx"),
        "jpg":  (_single_jpg,  ".jpg"),
        "png":  (_single_png,  ".png"),
    }
    if fmt not in dispatch:
        raise ValueError(f"Белгісіз формат: {fmt}")

    fn, ext = dispatch[fmt]
    return fn(img), filename + ext


def convert_multi(images_data: list[bytes], fmt: str, filename: str) -> tuple[bytes, str]:
    """Бірнеше суретті PDF немесе DOCX-қа біріктіру."""
    if fmt not in ("pdf", "docx"):
        raise ValueError("Көп сурет тек PDF / DOCX форматын қолдайды")

    imgs = [_open(d) for d in images_data]

    if fmt == "pdf":
        return _multi_pdf(imgs), filename + ".pdf"
    return _multi_docx(imgs), filename + ".docx"


FORMAT_LABELS = {
    "pdf":  "PDF",
    "docx": "Word (DOCX)",
    "jpg":  "JPG",
    "png":  "PNG",
}
